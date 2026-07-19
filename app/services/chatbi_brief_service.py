"""Evidence-backed, deterministic ChatBI business brief generation."""

from __future__ import annotations

import re
import tempfile
from pathlib import Path
from typing import Any, Iterable, Mapping

from docx import Document


class UnsupportedBriefClaim(ValueError):
    pass


def _rows(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [row for row in payload if isinstance(row, dict)]
    if isinstance(payload, Mapping):
        for key in ("rows", "data", "result", "items", "records"):
            if key in payload:
                nested = _rows(payload[key])
                if nested or payload[key] == []:
                    return nested
    return []


def _numeric_summary(rows: list[dict[str, Any]]) -> dict[str, dict[str, float]]:
    columns = {key for row in rows for key, value in row.items() if isinstance(value, (int, float)) and not isinstance(value, bool)}
    summary: dict[str, dict[str, float]] = {}
    for column in sorted(columns):
        values = [float(row[column]) for row in rows if isinstance(row.get(column), (int, float))]
        if values:
            summary[column] = {
                "sum": round(sum(values), 4),
                "average": round(sum(values) / len(values), 4),
                "min": min(values),
                "max": max(values),
            }
    return summary


def _validate_claims(claims: Iterable[str], supported_numbers: set[float]) -> None:
    for claim in claims:
        numbers = [float(item) for item in re.findall(r"-?\d+(?:\.\d+)?", str(claim))]
        if numbers and any(number not in supported_numbers for number in numbers):
            raise UnsupportedBriefClaim(f"简报主张缺少数据证据：{claim}")


def build_business_brief(
    result_ref: Mapping[str, Any],
    *,
    requested_claims: Iterable[str] = (),
) -> dict[str, Any]:
    rows = _rows(result_ref.get("rows"))
    numeric = _numeric_summary(rows)
    supported_numbers = {
        float(value)
        for row in rows
        for value in row.values()
        if isinstance(value, (int, float)) and not isinstance(value, bool)
    }
    for stats in numeric.values():
        supported_numbers.update(float(value) for value in stats.values())
    _validate_claims(requested_claims, supported_numbers)
    question = str(result_ref.get("question") or "数据分析").strip()
    title_core = re.sub(r"^(请|帮我|查询|查一下|查看)", "", question).strip() or "数据分析"
    title = f"{title_core}业务简报"
    context = dict(result_ref.get("analysis_context") or {})
    facts = {"row_count": len(rows), "numeric_summary": numeric, "analysis_context": context}
    lines = [
        f"# {title}",
        "",
        "## 分析范围",
        f"- 原始问题：{question}",
        f"- 数据集：{result_ref.get('dataset_name') or '未标注'}",
        f"- 数据行数：{len(rows)}",
        "",
        "## 核心数据",
    ]
    if numeric:
        for column, stats in numeric.items():
            lines.append(
                f"- {column}：合计 {stats['sum']:g}，平均 {stats['average']:g}，"
                f"最低 {stats['min']:g}，最高 {stats['max']:g}"
            )
    else:
        lines.append("- 本次结果不包含可汇总的数值字段，请以明细表为准。")
    lines.extend([
        "",
        "## 业务结论",
        f"- {str(result_ref.get('result_summary') or '已完成本次数据查询，结论仅基于上述结构化结果。')}",
        "",
        "## 后续动作",
        "- 复核异常值及业务口径。",
        "- 对关键指标设置周期性监控和阈值告警。",
        "",
        f"> 证据引用：ChatBI 结果 `{result_ref.get('result_id') or 'legacy'}`",
    ])
    return {
        "version": 1,
        "title": title,
        "facts": facts,
        "evidence": [{"result_id": result_ref.get("result_id"), "dataset_name": result_ref.get("dataset_name")}],
        "markdown": "\n".join(lines),
    }


def publish_business_brief_docx(brief: Mapping[str, Any]):
    """Create a private 24-hour Word artifact from deterministic brief sections."""
    from app.services.ai.tools.generated_file_service import publish

    safe_title = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", str(brief.get("title") or "业务简报"))[:60]
    with tempfile.TemporaryDirectory(prefix="chatbi_brief_") as temp_dir:
        path = Path(temp_dir) / f"{safe_title}.docx"
        document = Document()
        for line in str(brief.get("markdown") or "").splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("> "):
                document.add_paragraph(line[2:])
            elif line.strip():
                document.add_paragraph(line)
        document.save(path)
        return publish(path, path.name)
