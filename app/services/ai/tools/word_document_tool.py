"""Controlled Word document tools for configured agents."""
from __future__ import annotations

from typing import Any

from docx import Document

from app.core.context import get_current_agent_context
from app.services.ai.tools.document_paths import DocumentPathError, resolve_document_input_path, resolve_document_output_path
from app.services.ai.tools.generated_file_service import publish
from app.services.ai.tools.tool_compat import tool

_EXTENSIONS = {".docx"}


def _context():
    context = get_current_agent_context()
    if not context:
        raise DocumentPathError("无法获取当前执行上下文")
    return context


async def _input_path(path: str):
    context = _context()
    return await resolve_document_input_path(path, allowed_attachment_paths=context.authorized_attachment_paths, user_id=context.user_id, conversation_id=context.conversation_id, allowed_extensions=_EXTENSIONS)


async def _output_path(filename: str):
    context = _context()
    return await resolve_document_output_path(filename, user_id=context.user_id, conversation_id=context.conversation_id, allowed_extensions=_EXTENSIONS)


def _result(output_path, summary: str, changes: dict[str, Any]) -> dict[str, Any]:
    artifact = publish(output_path, output_path.name)
    return {"status": "ok", "summary": summary, "changes": changes, "artifact": artifact.to_tool_payload()}


@tool
async def word_document_read(action: str, path: str, start: int = 0, limit: int = 20) -> dict[str, Any]:
    """Inspect a Word document or read a bounded page of paragraphs."""
    if action not in {"inspect", "read_content"}:
        raise DocumentPathError("word_document_read 仅支持 inspect 或 read_content")
    document = Document(await _input_path(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    if action == "inspect":
        return {"status": "ok", "summary": f"文档包含 {len(paragraphs)} 个段落和 {len(document.tables)} 个表格", "data": {"paragraph_count": len(paragraphs), "table_count": len(document.tables), "preview": paragraphs[:20]}, "truncated": len(paragraphs) > 20}
    start = max(0, start)
    limit = min(max(1, limit), 50)
    page = paragraphs[start:start + limit]
    return {"status": "ok", "summary": f"已读取 {len(page)} 个段落", "data": {"paragraphs": page, "start": start}, "truncated": start + len(page) < len(paragraphs)}


@tool
async def word_document_write(action: str, output_filename: str, path: str | None = None, replacements: list[dict[str, str]] | None = None, paragraphs: list[str] | None = None, headers: list[str] | None = None, rows: list[list[str]] | None = None, title: str | None = None) -> dict[str, Any]:
    """Create or modify a Word document copy and return a download link.

    Copy artifact.download_url verbatim in the final response; never add a protocol or host.
    """
    if action not in {"create", "replace_text", "append_paragraphs", "append_table"}:
        raise DocumentPathError("word_document_write 不支持该操作")
    if action == "create":
        document = Document()
        if title:
            document.add_heading(title, level=1)
    else:
        if not path:
            raise DocumentPathError("修改文档需要 path")
        document = Document(await _input_path(path))
    changes: dict[str, Any] = {}
    if action == "replace_text":
        count = 0
        for replacement in replacements or []:
            find, replace = replacement.get("find", ""), replacement.get("replace", "")
            if not find:
                raise DocumentPathError("替换目标不能为空")
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
        changes["replacements"] = count
    elif action == "append_paragraphs":
        for text in paragraphs or []:
            document.add_paragraph(text)
        changes["appended_paragraphs"] = len(paragraphs or [])
    elif action == "append_table":
        if not headers:
            raise DocumentPathError("append_table 需要 headers")
        table = document.add_table(rows=1, cols=len(headers))
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)
        for row in rows or []:
            if len(row) != len(headers):
                raise DocumentPathError("表格行列数必须与 headers 一致")
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
        changes["appended_table_rows"] = len(rows or [])
    else:
        for text in paragraphs or []:
            document.add_paragraph(text)
        changes["created_document"] = True
    output_path = await _output_path(output_filename)
    document.save(output_path)
    return _result(output_path, "已生成 Word 文档", changes)
